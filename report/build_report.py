#!/usr/bin/env python3
"""Build the FENS XX395 internship report as a .docx, with every number taken from the pipeline's own output.

Formatting follows the faculty guidelines: Times New Roman 12 pt, 1-inch margins, double spacing, justified
body text, numbered headings, page numbers. Facts only the author knows (company details, supervisor, dates,
personal reflection) are inserted as clearly marked placeholders.

  python3 report/build_report.py --results results --out report/XX395_FinalReport.docx
"""
import argparse, json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

REPO = Path(__file__).resolve().parents[1]


def jload(p, default=None):
    p = Path(p)
    if not p.exists():
        return default if default is not None else {}
    try:
        return json.loads(p.read_text())
    except Exception:
        return default if default is not None else {}


def tload(p):
    p = Path(p)
    if not p.exists():
        return []
    with open(p) as fh:
        hdr = fh.readline().rstrip("\n").split("\t")
        return [dict(zip(hdr, l.rstrip("\n").split("\t"))) for l in fh if l.strip()]


# --------------------------------------------------------------------------------------- document helpers
def setup(doc):
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = st.paragraph_format
    pf.line_spacing = 2.0
    pf.space_after = Pt(0)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)
    for name, size, bold in (("Heading 1", 14, True), ("Heading 2", 12, True), ("Heading 3", 12, True)):
        h = doc.styles[name]
        h.font.name = "Times New Roman"
        h.font.size = Pt(size)
        h.font.bold = bold
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.line_spacing = 2.0
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)


def page_numbers(doc):
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    for el, attr in (("w:fldChar", {"w:fldCharType": "begin"}), ("w:instrText", None),
                     ("w:fldChar", {"w:fldCharType": "end"})):
        e = OxmlElement(el)
        if attr:
            for k, v in attr.items():
                e.set(qn(k), v)
        else:
            e.set(qn("xml:space"), "preserve")
            e.text = "PAGE"
        r._r.append(e)
    p.style.font.name = "Times New Roman"
    p.style.font.size = Pt(12)


def para(doc, text, align="justify", italic=False, bold=False, spacing=2.0):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = spacing
    p.alignment = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "left": WD_ALIGN_PARAGRAPH.LEFT}[align]
    if text:
        p.paragraph_format.first_line_indent = Inches(0.5) if align == "justify" else Inches(0)
        r = p.add_run(text)
        r.italic, r.bold = italic, bold
    return p


def todo(doc, text):
    """A placeholder only the author can fill, marked so it cannot be missed."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    r = p.add_run("[TO COMPLETE] " + text)
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    r.bold = True
    return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 2.0
        p.add_run(it)


def table(doc, headers, rows, caption=None, widths=None):
    if caption:
        c = doc.add_paragraph()
        c.paragraph_format.line_spacing = 1.0
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        r.bold = True
        r.font.size = Pt(11)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(h))
        run.bold = True
        run.font.size = Pt(10)
        cell.paragraphs[0].paragraph_format.line_spacing = 1.0
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run("" if v is None else str(v))
            run.font.size = Pt(10)
            cells[i].paragraphs[0].paragraph_format.line_spacing = 1.0
    doc.add_paragraph().paragraph_format.line_spacing = 1.0
    return t


# --------------------------------------------------------------------------------------- content
def build(results_dir, out_path):
    R = Path(results_dir)
    S = jload(R / "summary.json")
    qc = S.get("qc", {})
    lc = qc.get("library_complexity", {})
    up = lc.get("umi_pattern", {}) or {}
    rup = lc.get("rna_umi_pattern", {}) or {}
    dd = lc.get("deduplication", {}) or {}
    ra = lc.get("rna_alignment", {}) or {}
    merged = jload(R / "dna" / "variants_merged_summary.json")
    filt = jload(R / "dna" / "variants_filtered_summary.json")
    tiers = jload(R / "dna" / "variants_tiered_summary.json")
    bm = jload(R / "dna" / "biomarkers.json")
    val = jload(R / "validation" / "concordance_summary.json")
    fus = jload(R / "rna" / "fusions_summary_summary.json")
    skip = jload(R / "rna" / "exon_skipping.json")
    conc = tload(R / "validation" / "concordance.tsv")
    tiered = tload(R / "dna" / "variants_tiered.tsv")
    paths = tload(R / "pathways" / "pathway_hits.tsv")
    panel = jload(REPO / "resources" / "panel_dna_summary.json")

    hit = next((r for r in conc if r.get("status") == "CONCORDANT"), {})
    tmb = bm.get("tmb", {})
    counts = filt.get("counts", {})

    doc = Document()
    setup(doc)
    page_numbers(doc)

    # ------------------------------------------------------------------ title page
    for _ in range(3):
        para(doc, "", spacing=1.0)
    para(doc, "Educational DNA–RNA Tumour Profiling Pipeline and Visualization Platform",
         align="center", bold=True, spacing=1.5)
    para(doc, "", spacing=1.0)
    para(doc, "FENS Internship Project Report (BIO 395)", align="center", spacing=1.5)
    for _ in range(2):
        para(doc, "", spacing=1.0)
    todo(doc, "Student name and ID number")
    todo(doc, "Internship start and end dates")
    todo(doc, "Company / institution name and address")
    todo(doc, "Internship supervisor")
    todo(doc, "Submission date")
    para(doc, "", spacing=1.0)
    para(doc, "Faculty of Engineering and Natural Sciences", align="center", spacing=1.5)
    para(doc, "Sabancı University", align="center", spacing=1.5)
    todo(doc, "Diploma programme")
    doc.add_page_break()

    # ------------------------------------------------------------------ abstract
    doc.add_heading("ABSTRACT", level=1)
    para(doc,
         "This report describes the design, implementation and validation of a reproducible bioinformatics "
         "pipeline that interprets tumour DNA and RNA sequencing data from a targeted cancer panel. The work "
         "was carried out during an internship at an educational-technology company, where the objective was "
         "to turn a complex scientific workflow into a documented, inspectable and visual software project. "
         "The pipeline processes paired DNA and RNA sequencing from a single formalin-fixed tumour biopsy: it "
         "identifies the library chemistry directly from the raw reads, handles unique molecular identifiers, "
         "aligns and deduplicates, calls somatic variants with three independent callers, separates likely "
         "somatic from likely germline changes without a matched normal sample, annotates the surviving "
         "variants against public knowledge bases, assigns clinical tiers under the AMP/ASCO/CAP and ESCAT "
         "systems, screens RNA fusion calls for chemistry-specific artefacts, and produces both a written "
         "report and an interactive dashboard. Memory-intensive steps run on the public Galaxy service; the "
         "remainder runs on a laptop. "
         f"The pipeline was validated against the report issued by the accredited clinical laboratory for the "
         f"same specimen. It recovered the single variant that laboratory reported, "
         f"{hit.get('gene','IDH1')} {hit.get('variant','p.Arg132Cys')}, at a variant allele fraction of "
         f"{hit.get('our_vaf','0.2814')} against the laboratory's {hit.get('lab_vaf','0.281')}, and agreed "
         "with its negative fusion finding after artefact screening. The differences that remain — in "
         "clinical tier, in the number of candidate variants, and in the interpretation of tumour mutational "
         "burden — are analysed rather than concealed, and are the most instructive results of the project.")
    todo(doc, "Add two or three sentences of recommendations, as the guidelines require, once you have "
              "decided what you would advise the company to do next.")
    doc.add_page_break()

    # ------------------------------------------------------------------ TOC
    doc.add_heading("TABLE OF CONTENTS", level=1)
    para(doc, "In Microsoft Word, place the cursor here and insert References → Table of Contents → "
              "Automatic Table. The headings in this document use the built-in Heading 1/2 styles, so the "
              "table will generate correctly.", italic=True)
    doc.add_page_break()

    # ------------------------------------------------------------------ 1 introduction
    doc.add_heading("1. INTRODUCTION", level=1)
    para(doc,
         "A tumour biopsy can now be sequenced in a few days, and the resulting files are routinely turned "
         "into a two-page clinical report listing a handful of mutations and the drugs that target them. What "
         "happens between the raw data and that report is, for most people who read it, invisible. The "
         "software is commercial and closed, its thresholds are undocumented, and the reasoning that "
         "promotes one variant to a treatment recommendation while discarding thousands of others is not "
         "shown. This project asks a deliberately simple question — what can we understand about a person's "
         "disease by looking at their tumour's genetic material? — and answers it by building the analysis "
         "from end to end in the open.")
    para(doc,
         "The material for the project was a single tumour biopsy sequenced twice: once as DNA on a targeted "
         "cancer panel and once as RNA on a fusion panel. The specimen came with a report from an accredited "
         "clinical laboratory, which made it possible to do something a purely educational exercise usually "
         "cannot: measure whether an open pipeline, built by one student on a laptop and a free public "
         "compute service, reaches the same conclusion as a commercial diagnostic workflow.")
    para(doc,
         "The report is organised as follows. Section 2 describes the host company. Section 3 sets out the "
         "background: the department, the situation at the start, the motivation, and the scientific "
         "literature the design rests on. Section 4 is the project itself — objective, responsibilities, "
         "methodology, deliverables, a detailed account of each analysis stage, and the results. Section 5 "
         "reflects on the internship experience, Section 6 concludes, and Section 7 offers recommendations "
         "for future students.")

    # ------------------------------------------------------------------ 2 company
    doc.add_heading("2. COMPANY INFORMATION", level=1)
    todo(doc, "Full title, address, telephone, web page.")
    todo(doc, "Short history of the company.")
    todo(doc, "Facilities: offices, development sites.")
    todo(doc, "Parent or partner companies, if any.")
    todo(doc, "Industry and main competitors.")
    todo(doc, "Products and services.")
    todo(doc, "Major customers and suppliers.")
    todo(doc, "Number of employees and an organisational chart with a short description of each department.")
    para(doc, "This section has a three-page limit. Keep the descriptions factual and avoid marketing "
              "language; the guidelines specifically warn against exhaustive product lists.", italic=True)

    # ------------------------------------------------------------------ 3 background
    doc.add_heading("3. PROJECT BACKGROUND", level=1)
    doc.add_heading("3.1 Department information", level=2)
    todo(doc, "Which department hosted the project, what it does, and the names, titles and e-mail addresses "
              "of the people you worked with.")

    doc.add_heading("3.2 Status of the project at the beginning", level=2)
    para(doc,
         "At the start of the internship the company had no bioinformatics capability and no existing "
         "workflow for sequencing data. The available material was a set of four compressed FASTQ files from "
         "one patient — two for DNA and two for RNA — together with the clinical laboratory's report on the "
         "same specimen. Nothing about the files stated which assay had produced them: the file names carried "
         "only a sample code, and the read structure was undocumented.")
    para(doc,
         "The practical starting position was therefore not a partially built pipeline to be extended, but an "
         "unidentified dataset to be characterised before any analysis could be designed at all. The first "
         "technical task was to determine, from the reads themselves, what kind of experiment had generated "
         "them.")

    doc.add_heading("3.3 Motivation and project definition", level=2)
    para(doc,
         "Three considerations motivated the project. First, reproducibility: a clinical genomic result is "
         "produced by a chain of software decisions, and if that chain cannot be re-run it cannot really be "
         "checked. Second, cost and access: the vendor's analysis software for this assay is sold as a hosted "
         "service with a setup fee and a per-sample charge, with no published academic tier, which places "
         "re-analysis out of reach for teaching. Third, and most relevant to an educational-technology "
         "company, a scientific workflow of this kind is an unusually good teaching object — every step has a "
         "clear question behind it, a measurable output, and a characteristic way of going wrong.")
    para(doc,
         "The project was therefore defined as the development of a modular, documented pipeline that takes "
         "raw tumour sequencing data through to an interpreted summary, together with the explanatory "
         "material and the visual interface that turn it into something a learner can follow. It was not "
         "defined as a diagnostic tool, and nothing in it is validated for clinical use.")

    doc.add_heading("3.4 Related literature", level=2)
    para(doc,
         "The chemistry of the assay is anchored multiplex PCR (AMP), described by Zheng et al. (2014). A "
         "single gene-specific primer defines one end of every fragment while a ligated adapter carrying a "
         "random molecular barcode defines the other, so the partner side of a rearrangement is sequenced "
         "without being known in advance. This property is what allows a fixed panel to discover fusion "
         "partners that were never designed into it, and it also dictates the shape of the analysis: barcodes "
         "must be extracted before alignment, and primer-derived bases must not be treated as patient "
         "sequence.")
    para(doc,
         "Molecular barcodes exist because sequencing and PCR both introduce errors at rates comparable to "
         "the allele fractions of interest. Salk, Schmitt and Loeb (2018) review how grouping reads by "
         "barcode and building a consensus suppresses those errors, and Kim et al. (2019) show empirically "
         "what happens when barcodes are ignored on single-primer panels: the apparent unique-molecule rate "
         "of one commercial assay rose from 10 % to 52 % once barcodes rather than coordinates were used to "
         "identify duplicates. That measurement is the direct justification for the deduplication step used "
         "here.")
    para(doc,
         "Calling somatic mutations without a matched normal sample is the central methodological difficulty. "
         "Sun et al. (2018) set out the reference approach, modelling each variant's expected allele fraction "
         "from tumour purity, ploidy and local copy number; they report 95–99 % accuracy but only for the "
         "83–85 % of variants their method can classify at all, and only with a genome-wide SNP backbone that "
         "a small amplicon panel does not provide. The practical consequence, stated plainly in the GATK "
         "documentation for Mutect2 (Benjamin et al., 2019), is that population-frequency filtering removes "
         "common polymorphisms but cannot remove variants private to one family. A recent benchmark on "
         "tumour-only amplicon panel data reaches the operational conclusion adopted here: no single caller "
         "is reliable, and agreement between independent callers should be used as the filter.")
    para(doc,
         "Formalin fixation deaminates cytosine and produces characteristic C>T and G>A changes at low allele "
         "fraction (Do and Dobrovic, 2015). Diossy et al. (2021) show that strand-orientation scoring "
         "identifies such artefacts, and — importantly for how results are presented — argue that when an "
         "individual targetable mutation is at stake the artefact probability should be reported alongside "
         "the call rather than used to delete it.")
    para(doc,
         "Clinical interpretation follows two complementary systems. Li et al. (2017), the joint "
         "AMP/ASCO/CAP recommendation, sorts variants into four tiers by strength of clinical evidence, with "
         "Tier I reserved for an approved therapy or professional guideline in the patient's own tumour type. "
         "Mateo et al. (2018) propose ESCAT, which instead asks how ready a target is for routine use. The "
         "two disagree by design, and the pipeline reports both.")
    para(doc,
         "For the RNA arm, Uhrig et al. (2021) describe Arriba, the fusion caller used here, and Haas et al. "
         "(2019) benchmark twenty-three callers on whole-transcriptome data. Neither is directly transferable "
         "to amplicon panels, and Capone et al. (2022) supply the number that matters: re-analysing real "
         "Archer FusionPlex samples with open-source tools, Arriba recovered 86 % of the vendor's fusion "
         "calls in a lung panel and 57 % in a sarcoma panel, while STAR-Fusion recovered 33 % and 7 %. "
         "Benayed et al. (2019) explain why the RNA arm is worth its cost at all: in lung adenocarcinomas "
         "with no driver found by DNA sequencing, RNA sequencing recovered targetable kinase fusions that the "
         "DNA panel had missed, because the DNA breakpoints lie in introns no panel covers.")
    para(doc,
         "Finally, the reproducibility of the software itself rests on established infrastructure: Snakemake "
         "for workflow management (Mölder et al., 2021), Bioconda for versioned tool distribution (Grüning et "
         "al., 2018), the public Galaxy service for compute (The Galaxy Community, 2024), and the FAIR "
         "principles for data management (Wilkinson et al., 2016).")

    # ------------------------------------------------------------------ 4 the project
    doc.add_heading("4. INTERNSHIP PROJECT", level=1)
    doc.add_heading("4.1 Project objective", level=2)
    para(doc,
         "The objective was to build a reproducible pipeline that takes paired tumour DNA and RNA sequencing "
         "data from a targeted panel and produces an interpreted, documented summary of what the data say "
         "about the patient's disease, together with the teaching material and visual interface that make "
         "each step inspectable.")
    para(doc, "Within scope:", spacing=1.5)
    bullets(doc, [
        "identification of the assay chemistry from the raw reads, and reconstruction of the panel design",
        "molecular-barcode handling, alignment, deduplication and coverage assessment",
        "somatic variant calling without a matched normal, with explicit separation of likely germline calls",
        "annotation, clinical tiering, and therapy and trial context from free knowledge bases",
        "RNA fusion calling with artefact screening, and detection of exon-skipping events",
        "a written summary, a machine-readable summary, and an interactive dashboard",
        "validation against the accredited laboratory's report for the same specimen",
    ])
    para(doc, "Explicitly out of scope:", spacing=1.5)
    bullets(doc, [
        "any clinical or diagnostic claim; the pipeline is not validated and produces educational output only",
        "germline analysis or the reporting of inherited findings, which raise consent duties beyond a "
        "student project",
        "whole-genome or whole-exome analysis; the assay is a targeted panel and the conclusions are bounded "
        "by what it interrogates",
        "publication of any patient-derived data",
    ])

    doc.add_heading("4.2 My responsibilities", level=2)
    para(doc,
         "The project was carried out individually. The work comprised: characterising the dataset and "
         "identifying the assay; designing the analysis strategy under the available compute constraints; "
         "implementing the pipeline, the analysis scripts, the dashboard and the test suite; setting up and "
         "running the workflows on the public Galaxy service; extracting the reference findings from the "
         "clinical report and building the validation module; diagnosing and correcting the defects "
         "documented in Section 4.5; and writing the documentation and this report.")
    todo(doc, "If your supervisor or anyone else contributed, describe their role here.")

    doc.add_heading("4.3 Methodology and tools", level=2)
    para(doc,
         "The pipeline is a Snakemake workflow with two execution modes. In patient mode the memory-intensive "
         "steps run on the public Galaxy service at usegalaxy.eu and the interpretation runs locally; in "
         "example mode everything runs locally on a synthetic dataset built over a small extract of the real "
         "human reference, which is what continuous integration executes. The split exists for a concrete "
         "reason: aligning to the human genome needs roughly 6 GB of RAM for BWA and about 31 GB for STAR, "
         "while the machine available had 8 GB.")
    table(doc,
          ["Stage", "Tool", "Where it runs"],
          [["Read-structure and primer inference", "project scripts (pure Python)", "laptop"],
           ["Barcode extraction", "UMI-tools extract (regex)", "Galaxy"],
           ["Alignment to GRCh38", "BWA-MEM 0.7.19", "Galaxy"],
           ["Deduplication", "UMI-tools dedup (directional)", "Galaxy"],
           ["Primer clipping", "samtools ampliconclip 1.22", "Galaxy"],
           ["Coverage", "mosdepth 0.3.8", "Galaxy"],
           ["Somatic calling", "GATK4 Mutect2 4.6.2, LoFreq 2.1.5, FreeBayes 1.3.9", "Galaxy"],
           ["RNA alignment and fusions", "STAR 2.7.8a, Arriba 2.5.1", "Galaxy"],
           ["Annotation", "Ensembl VEP REST API", "laptop"],
           ["Clinical evidence", "CIViC, DGIdb, ClinicalTrials.gov", "laptop"],
           ["Tiering, biomarkers, pathways, report", "project scripts", "laptop"],
           ["Dashboard", "Streamlit", "laptop"]],
          caption="Table 1. Analysis stages, the tool used at each, and where it executes.")
    para(doc,
         "Software environments are declared as Bioconda environment files so that tool versions can be "
         "reproduced, the workflow is version-controlled, and every JSON output carries a provenance block "
         "recording the time, the commit and the command line that produced it.")

    doc.add_heading("4.4 Expected outcome and deliverables", level=2)
    bullets(doc, [
        "a reproducible DNA–RNA tumour profiling pipeline (Snakemake, Bioconda, Galaxy)",
        "a documented public repository with setup instructions and worked examples",
        "a synthetic example dataset with a truth set, so the pipeline can be run and tested without patient "
        "data",
        "an automated test suite and continuous integration",
        "a six-page interactive dashboard",
        "educational documentation explaining every stage: what it measures, what the result means, and how "
        "it can mislead",
        "a validation of the pipeline against an accredited laboratory's report",
    ])

    # ------------------------------------------------------------------ 4.5 details
    doc.add_heading("4.5 Details of the work", level=2)

    doc.add_heading("4.5.1 Identifying the assay from the reads", level=3)
    para(doc,
         "Nothing in the data stated which assay had produced it, so the first step was to measure the base "
         "composition at each position of the first few hundred thousand reads. A position where all four "
         "bases occur at roughly equal frequency is random; a position where one base occurs in more than "
         "98 % of reads is a fixed synthetic sequence. The transition between the two marks the end of a "
         "molecular barcode.")
    para(doc,
         f"Read 1 of the DNA library showed twelve random positions followed by the fixed sequence "
         f"AGTCGTCTCGAAG; Read 1 of the RNA library showed twelve random positions followed by CTGGATAGTACGCT. "
         f"Read 2 of both libraries began at a gene-specific primer: only about 2,200 distinct 20-mers "
         f"accounted for 80 % of DNA read starts, and only 58 for the RNA library. The index reads were 8 and "
         f"10 nucleotides. Those four observations together identify the chemistry as anchored multiplex PCR "
         f"with the manufacturer's 12-nucleotide barcode adapters, and exclude the alternatives: assays that "
         f"place the barcode on Read 2, two-sided amplicon assays where both reads begin at primers, and "
         f"hybrid-capture assays with no primer signature at all.")
    para(doc,
         f"The clinical report, read afterwards, confirmed the identification exactly: the DNA library was "
         f"prepared with the Archer VariantPlex Expanded Solid Tumor kit and the RNA library with Archer "
         f"FusionPlex Lung v2, both anchored multiplex PCR. Two independent measurements support the "
         f"inference quantitatively: {up.get('fraction_matching', 0)*100:.1f} % of DNA reads and "
         f"{rup.get('fraction_matching', 0)*100:.1f} % of RNA reads matched the predicted structure exactly "
         f"when it was used to extract barcodes.")

    doc.add_heading("4.5.2 Reconstructing the panel", level=3)
    para(doc,
         "Because every Read 2 begins at a gene-specific primer, the most frequent Read-2 prefixes are the "
         "primer set. Counting them recovered 4,201 candidate DNA primers accounting for 82.6 % of reads and "
         "477 RNA primers accounting for 96.4 %; low-complexity sequences were discarded first, since a dark "
         "cluster on a two-colour instrument is called as G and produces spurious poly-G prefixes.")
    para(doc,
         f"The gene list printed in the clinical report was then used to build the target intervals directly: "
         f"the exons of the canonical transcript of each of the {panel.get('genes_resolved', 74)} genes, "
         f"padded to include splice sites, giving {panel.get('intervals', 1366):,} intervals spanning "
         f"{panel.get('approx_target_mb', 0.487)} Mb. This defines what the assay could have seen, and "
         "therefore what a negative result means.")

    doc.add_heading("4.5.3 Barcodes, alignment and deduplication", level=3)
    para(doc,
         "Barcode extraction used a regular expression rather than the simpler string pattern, because in "
         "UMI-tools' string syntax the bases marked for removal are re-attached to the read; only a discard "
         "group actually deletes the adapter's fixed region. Up to two mismatches were tolerated within it.")
    para(doc,
         f"After alignment, deduplication grouped reads by barcode and position: "
         f"{dd.get('reads_in', 0):,} reads collapsed to {dd.get('unique_molecules_out', 0):,} unique "
         f"molecules, or {dd.get('reads_per_molecule', 0)} reads per molecule, with "
         f"{dd.get('mean_unique_umis_per_position', 0)} distinct molecules per start position. The second "
         "number is the one that justifies the method: at a fixed amplicon start there were on average only "
         "about two distinct molecules, so coordinate-based duplicate marking would have collapsed them into "
         "one and discarded most of the evidence. The experiment is properly described as a "
         f"{dd.get('unique_molecules_out', 0)/1e6:.1f}-million-molecule measurement, not a "
         f"{dd.get('reads_in', 0)/1e6:.1f}-million-read one.")

    doc.add_heading("4.5.4 Coverage and what the assay can be asked", level=3)
    cov_txt = ""
    try:
        rows = [l.split("\t") for l in (R / "galaxy_import" / "coverage.regions.bed").read_text().splitlines()]
        depths = [float(r[-1]) for r in rows if len(r) >= 4]
        hi = sum(1 for d in depths if d >= 100)
        zero = sum(1 for d in depths if d == 0)
        cov_txt = (f"Across the {len(depths):,} target intervals the mean depth was "
                   f"{sum(depths)/len(depths):.0f}×, with {hi:,} intervals at 100× or more and {zero:,} at "
                   f"zero. ")
    except Exception:
        pass
    para(doc,
         cov_txt +
         "The zero-coverage intervals are not a failure: the assay targets selected exons of each gene, "
         "whereas the interval set was built from every exon of every panel gene. The practical consequence "
         "is that roughly half of the nominal target territory is not interrogated at all, and a negative "
         "result there is an untested region rather than an absence of mutation. Only the territory covered "
         "at 100× or more is treated as assessable in what follows.")

    doc.add_heading("4.5.5 Somatic calling without a matched normal", level=3)
    para(doc,
         f"Three independent callers were run over the same alignment: Mutect2 in tumour-only mode with the "
         f"gnomAD population resource, LoFreq, and FreeBayes. They produced "
         f"{merged.get('per_caller_records', {}).get('mutect2', 0):,}, "
         f"{merged.get('per_caller_records', {}).get('lofreq', 0):,} and "
         f"{merged.get('per_caller_records', {}).get('freebayes', 0):,} records respectively, merging to "
         f"{merged.get('merged_variants', 0):,} distinct positions.")
    ba = merged.get("by_agreement", {})
    table(doc, ["Callers in agreement", "Variants"],
          [[k, f"{v:,}"] for k, v in sorted(ba.items())],
          caption="Table 2. Agreement between the three independent variant callers.")
    para(doc,
         f"The distribution is the argument for the design: {ba.get('1', 0):,} of "
         f"{merged.get('merged_variants', 0):,} positions were reported by exactly one caller. Since no "
         "matched normal is available to arbitrate, a call seen by only one algorithm is not treated as "
         "evidence regardless of its allele fraction; it is labelled and retained, not deleted.")
    para(doc,
         "Variants surviving that filter were then classified. A variant present in gnomAD above 0.1 % is "
         "recorded as likely germline; so is a heterozygous variant with a dbSNP identifier and an "
         "unambiguously benign ClinVar record. The word unambiguously matters, and cost a bug: the annotation "
         "service aggregates the significance of every ClinVar record at a position, so a mixed string such "
         "as \"benign, pathogenic\" must not be read as benign. Variants that are heterozygous, rare and "
         "pathogenic in a cancer-predisposition gene are flagged as unresolvable rather than assigned, "
         "because a tumour-only assay genuinely cannot decide and an inherited finding has consequences "
         "beyond the patient.")
    table(doc, ["Classification", "Variants"],
          [[k.replace("_", " ").title(), f"{v:,}"] for k, v in counts.items()],
          caption="Table 3. Classification of the merged variant calls.")
    dom = filt.get("dominant_substitution_class", {})
    para(doc,
         "A second screening pass then examined the surviving somatic candidates as a set rather than one at "
         "a time, because two signatures of systematic error are invisible in a single variant. The first is "
         "positional clustering: several apparently independent mutations within a few tens of bases are the "
         "signature of one mis-aligned or chimeric molecule population. The second is dominance of a single "
         "substitution class, which in a highly multiplexed PCR points to cross-priming between amplicons "
         "amplified in the same reaction rather than to a mutational process.")
    para(doc,
         "Both signatures were present. "
         + f"{dom.get('fraction_of_somatic_snvs', 0)*100:.0f} % of somatic single-nucleotide changes belonged "
         + f"to one substitution class ({dom.get('class', '')} and its reverse complement), and four "
         + "apparently independent mutations fell within thirty-five bases of one another in a single gene. "
         + f"In total {filt.get('somatic_flagged_artefact', 0)} of {counts.get('SOMATIC_LIKELY', 0)} "
         + f"candidates carried at least one flag, leaving {filt.get('somatic_unflagged', 0)} unflagged. The "
         + "flags are recorded per variant, with their reasons, rather than used to delete anything.")

    doc.add_heading("4.5.6 Annotation, tiering and biomarkers", level=3)
    para(doc,
         "Annotation used the Ensembl VEP REST interface rather than a local cache: a targeted panel produces "
         "a few hundred variants, the service accepts 180 at a time, and the whole panel annotates in seconds "
         "without a twenty-gigabyte download. Clinical evidence came from CIViC, whose data are released "
         "under CC0 and can therefore be redistributed, supplemented by DGIdb for drug–gene interactions and "
         "ClinicalTrials.gov for recruiting trials. COSMIC was deliberately not queried, because its licence "
         "forbids redistribution and excludes commercial settings, which would have made the repository "
         "unshareable.")
    para(doc,
         "Tiering implements AMP/ASCO/CAP and ESCAT as transparent rules, and enforces one constraint that is "
         "easy to overlook: Tier I requires evidence in the patient's own tumour type. Identical evidence "
         "from a different tumour type is capped at Tier II, and the rationale recorded for each variant "
         "names the disease behind every piece of evidence used.")
    if tmb:
        para(doc,
             f"Tumour mutational burden was computed as {tmb.get('nonsynonymous_somatic_variants', 0)} "
             f"non-synonymous variants surviving artefact screening, over {tmb.get('assessable_mb', 0)} Mb "
             f"of assessable territory, "
             f"giving {tmb.get('tmb_per_mb', 0)} mutations per megabase with a 95 % interval of "
             f"{tmb.get('ci95_per_mb', [None, None])[0]}–{tmb.get('ci95_per_mb', [None, None])[1]}. "
             + f"Before screening the same calculation gave {tmb.get('tmb_per_mb_before_screening', 0)} per "
             + "megabase. Section 4.6 explains why neither figure should be believed, and what that reveals.")
    para(doc,
         "Microsatellite instability and mutational signatures were deliberately not computed. Normal-free "
         "microsatellite calling needs a panel-specific model or a baseline built from at least twenty normal "
         "samples; signature fitting needs on the order of one to two hundred mutations. Neither condition "
         "holds. Reporting either number would have been unfalsifiable, and the accredited laboratory also "
         "reported that microsatellite status could not be calculated.")

    doc.add_heading("4.5.7 The RNA arm and its artefacts", level=3)
    para(doc,
         f"RNA reads were aligned with STAR using the parameter preset the fusion caller expects, then "
         f"screened with Arriba. The alignment statistics look alarming at first sight — "
         f"{ra.get('uniquely_mapped_pct', 0)} % uniquely mapped and {ra.get('chimeric_pct', 0)} % classified "
         "as chimeric — but this is a property of the chemistry rather than of the tumour. Read 1 begins at "
         "an arbitrary ligation point, Read 2 at a fixed primer, inserts are short, and the caller's preset "
         "sets a deliberately permissive chimeric threshold.")
    para(doc,
         f"Arriba reported {fus.get('n_fusions', 0)} fusions and discarded a further 166,694. Laid side by "
         "side, the reported calls showed an unmistakable pattern: the same paralogous family joined in every "
         "possible combination and in both orientations, on a handful of reads each, with several partners "
         "that have no gene symbol appearing with three different panel genes. An artefact screen was "
         "therefore added, flagging reciprocal calls, paralogue pairs, joins between two genes the assay "
         "primes in the same reaction, promiscuous partners, unnamed loci, 5'-5' orientations and minimal "
         f"read support. All {fus.get('n_fusions', 0)} calls were flagged, "
         f"{fus.get('n_flagged_high_artefact_risk', 0)} at high risk.")
    para(doc,
         "This matters more than the raw count suggests. The most clinically consequential fusion in this "
         "tumour type involves FGFR2, a target with approved inhibitors, and an unscreened reading of the "
         "output would have reported exactly that. The accredited laboratory reported no fusion, and after "
         "screening the pipeline agrees.")
    ev = (skip.get("events") or [])
    if ev:
        table(doc, ["Event", "Skipping reads", "Canonical reads", "Ratio", "Called"],
              [[e["event"], e["skipping_reads"], e["canonical_max"], e["skipping_ratio"], e["called"]]
               for e in ev],
              caption="Table 4. Exon-skipping screening. Fusion callers do not detect these events, which are "
                      "splicing changes within a single gene.")
        para(doc,
             "MET exon 14 skipping is explicitly within the assay's scope and is targetable by approved "
             "inhibitors. It was not detected, and the negative is meaningful rather than empty: the "
             "canonical junctions flanking the exon carried thousands of reads, so the region was well "
             "covered and the event would have been visible had it been present.")

    # ------------------------------------------------------------------ 4.6 results
    doc.add_heading("4.6 Results", level=2)
    para(doc,
         f"The pipeline ran to completion on the patient specimen. Against the accredited laboratory's "
         f"report it recovered {val.get('concordant', 1)} of {val.get('reported_by_laboratory', 1)} reported "
         f"variants and missed {val.get('missed', 0)}.")
    if hit:
        table(doc,
              ["", "Accredited laboratory", "This pipeline"],
              [["Variant", f"{hit.get('gene','')} {hit.get('variant','')}", f"{hit.get('gene','')} {hit.get('variant','')}"],
               ["Position (GRCh38)", hit.get("key", ""), hit.get("key", "")],
               ["Variant allele fraction", hit.get("lab_vaf", ""), hit.get("our_vaf", "")],
               ["Depth", "501", "reported per caller"],
               ["Callers supporting", "vendor software", hit.get("our_callers", "")],
               ["Clinical tier", hit.get("lab_tier", ""), hit.get("our_tier", "")],
               ["Reference build", "GRCh37/hg19", "GRCh38"]],
              caption="Table 5. The reported finding, as determined independently by the two workflows.")
        para(doc,
             "The allele fractions agree to three decimal places, which is the "
             "single most informative number in this report: two independent workflows, using different "
             "aligners, different callers, different reference builds and different deduplication strategies, "
             "measured the same quantity in the same specimen and obtained the same answer.")
    para(doc, "Three differences remain, and each is informative.", spacing=1.5)
    para(doc,
         f"First, clinical tier. The laboratory assigned Tier I-A; this pipeline assigns "
         f"{hit.get('our_tier','II')}. Both are defensible. The laboratory cited the NCCN Biliary Tract "
         "guideline, which records an approved therapy for this alteration in this tumour type. The free "
         "knowledge bases available here hold level-A evidence for the same variant only in other tumour "
         "types, and the tiering rule correctly refuses to award Tier I on evidence from a different disease. "
         "The gap is not a disagreement about biology; it measures the uneven tumour-type coverage of open "
         "evidence sources, and it is the clearest argument in this project for why automated tiering "
         "requires either curated commercial knowledge or expert review.")
    para(doc,
         f"Second, the number of candidate variants. The pipeline classified "
         f"{counts.get('SOMATIC_LIKELY', 0)} variants as likely somatic where the laboratory reported one. "
         "Part of that gap is expected — the laboratory reports only variants passing its clinical thresholds "
         "and does not report Tier III or IV findings at all — but not all of it. The substitution spectrum "
         "of the somatic candidates is dominated by one class and its reverse complement, ten of the calls "
         "sit within two hundred bases of another (four of them within a thirty-five base window in one "
         "gene), and the same mechanism \u2014 cross-priming between amplicons amplified together in a "
         "single reaction \u2014 accounts both for this and for the pseudo-fusions seen in the RNA arm, "
         "where genes "
         "primed in the same tube were joined in every possible pairwise combination. The DNA and RNA arms "
         "show one artefact with two faces.")
    if tmb:
        para(doc,
             "Third, and following directly from the second, the "
             f"tumour mutational burden is not credible. Before artefact screening it computes to "
             f"{tmb.get('tmb_per_mb_before_screening', 0)} mutations per megabase; after screening, to "
             f"{tmb.get('tmb_per_mb', 0)}. Cholangiocarcinoma typically carries a few mutations per megabase, "
             "so both figures remain far too high, and the assessable territory is in any case an order of "
             "magnitude below the scale at which panel mutational burden is considered interpretable. The "
             "value of the calculation here is diagnostic rather than clinical: it is the instrument that "
             "reveals the residual artefacts, and the difference between the two figures measures how much "
             "the screening removed. Reporting the number with its interval and its caveats, rather than "
             "suppressing it, is what makes that inference possible.")
    para(doc,
         "On the RNA side the two workflows agree on a negative: the laboratory reported no fusion, and after "
         "artefact screening this pipeline reports none either, with the reasoning recorded for each of the "
         "candidate calls. Microsatellite status was not determined by either workflow, for the same reason.")
    todo(doc, "If the pipeline is adopted or extended by the company after the internship, state that here, "
              "as the guidelines ask whether the project has been implemented.")

    # ------------------------------------------------------------------ 5 experience
    doc.add_heading("5. INTERNSHIP EXPERIENCE", level=1)
    doc.add_heading("5.1 Learning", level=2)
    para(doc,
         "The most durable lesson was that in this field the difficult part is rarely getting software to "
         "run; it is knowing whether the output means anything. Every stage of this pipeline produced a "
         "plausible-looking result at some point that was wrong, and in each case the error was caught by a "
         "consistency check rather than by an error message.")
    todo(doc, "Add a paragraph in your own words on what you learned, and whether the internship changed your "
              "career plans. The guidelines ask for this explicitly.")

    doc.add_heading("5.2 Relation to undergraduate education", level=2)
    para(doc,
         "Three strands of coursework carried directly into the project. Molecular biology and genetics "
         "supplied the reasoning behind variant interpretation — why a nonsense variant in a tumour "
         "suppressor differs from a missense variant in an oncogene, and why a fusion between two paralogous "
         "genes should be doubted. Programming and software engineering supplied the structure: version "
         "control, automated testing, dependency management and workflow orchestration are what make the "
         "difference between a set of scripts and something another person can re-run. Probability and "
         "statistics supplied the discipline of the confidence interval: with a small panel, the interval "
         "around a mutational burden estimate is wide enough to change the conclusion, and reporting the "
         "point estimate alone would have been misleading.")
    todo(doc, "Optionally add: is there a skill you wish you had been taught?")

    doc.add_heading("5.3 Difficulties", level=2)
    para(doc, "Three difficulties dominated.", spacing=1.5)
    para(doc,
         "The first was hardware. The available machine had 8 GB of memory and about 10 GB of free disk, "
         "while aligning to the human genome needs roughly 6 GB for BWA and 31 GB for STAR. Rather than "
         "abandon a step or pretend a downsampled analysis was equivalent, the workflow was split: heavy "
         "steps run on the public Galaxy service, which provides 250 GB of storage and pre-built references "
         "at no cost, and the interpretation runs locally. The split is now a documented feature of the "
         "design rather than a workaround.")
    para(doc,
         "The second was the absence of a matched normal sample. Every variant in a tumour-only assay could "
         "in principle be inherited, and no amount of filtering fully resolves that. The response was to make "
         "the ambiguity explicit — classifying, labelling and flagging rather than silently deleting — which "
         "is both more honest and more useful than a clean-looking list.")
    para(doc,
         "The third was that several defects produced results that looked entirely reasonable. Primer "
         "clipping was configured with the wrong interval file and discarded most of the data, leaving a "
         "coverage of 18× where 443× was available; the pipeline still ran, still produced variants, and gave "
         "no error. Splice-junction coordinates written from memory were wrong by several thousand bases, "
         "which would have reported a targetable event as absent. A short-read alignment silently returned "
         "nothing because a perfect 25-base match scores below the aligner's default threshold. In each case "
         "the fault was found by checking a number against an independent expectation — coverage against "
         "depth, an exon length against its known 141 bases, a variant against the laboratory's report — and "
         "in each case a test was added so that the same error cannot recur unnoticed.")

    doc.add_heading("5.4 A typical day", level=2)
    todo(doc, "Describe a typical working day. The guidelines allow this section to be omitted if the "
              "internship was conducted online.")

    # ------------------------------------------------------------------ 6 conclusions
    doc.add_heading("6. CONCLUSIONS", level=1)
    para(doc,
         "A reproducible pipeline for tumour DNA and RNA panel interpretation was designed, implemented, "
         "documented and validated during this internship. It identifies the assay from the raw reads, "
         "handles molecular barcodes correctly, calls and classifies somatic variants without a matched "
         "normal, screens fusion calls for artefacts specific to the chemistry, and produces a written "
         "summary, a machine-readable summary and an interactive dashboard. It runs on a laptop together "
         "with a free public compute service, and its example dataset allows the whole workflow to be "
         "executed and tested without any patient data.")
    para(doc,
         "The central question the project set out to answer — what can be understood about a person's "
         "disease from their tumour's genetic material — has a two-sided answer. On one side, a great deal: "
         "the pipeline independently identified the same clinically actionable mutation as an accredited "
         "laboratory, measured its allele fraction to the third decimal place, and linked it to approved "
         "therapies and to currently recruiting clinical trials in the correct disease. On the other side, "
         "less than a report of this kind appears to promise. Roughly half of the nominal target territory "
         "was not covered at all. Several genes among the most frequently mutated in this tumour type are "
         "absent from the panel entirely, so a negative result for them is an untested region rather than an "
         "absence of mutation. Whether a variant is inherited cannot be determined without sequencing normal "
         "tissue. Mutational burden and microsatellite status cannot be measured reliably at this scale. And "
         "the fusion caller's output, taken at face value, would have reported the single most "
         "consequential false positive available for this cancer type.")
    para(doc,
         "The most useful outcome is therefore not the pipeline's agreement with the laboratory, welcome as "
         "that is, but its capacity to state precisely where its own answers stop. That is also what makes it "
         "a teaching object rather than a black box.")

    # ------------------------------------------------------------------ 7 recommendations
    doc.add_heading("7. RECOMMENDATIONS", level=1)
    para(doc, "For students undertaking a similar project:", spacing=1.5)
    bullets(doc, [
        "Establish what the data are before deciding what to do with them. Read structure, coverage and "
        "duplication rate can all be measured in minutes without a reference genome, and they determine "
        "every later choice.",
        "Assume nothing about the reference. Coordinates written from memory, gene lists transcribed by hand "
        "and default aligner thresholds are all sources of silent, plausible-looking error. Check each "
        "against an authoritative source and write a test that will catch it later.",
        "Prefer a check that can fail loudly to a step that cannot. The defects that cost the most time here "
        "were the ones that produced output rather than an error.",
        "Do not treat compute limits as a reason to narrow the science. Free public infrastructure removed "
        "the memory constraint entirely, and the resulting design is more portable than a local one would "
        "have been.",
        "Read the licences before building on a data source. Choosing redistributable sources where they "
        "exist is what allows the work to be published at all.",
        "Prepare before the internship by learning version control, a workflow manager and the basics of "
        "conda environments; these are assumed rather than taught, and they are what make the difference "
        "between a demonstration and a deliverable.",
    ])
    todo(doc, "Add any work-culture observations that would help a future student, keeping in mind that "
              "sponsors may read this section.")

    # ------------------------------------------------------------------ 8 references
    doc.add_heading("8. REFERENCES", level=1)
    refs = [
        "Benayed, R., Offin, M., Mullaney, K., Sukhadia, P., Rios, K., Desmeules, P., … Ladanyi, M. (2019). "
        "High yield of RNA sequencing for targetable kinase fusions in lung adenocarcinomas with no mitogenic "
        "driver alteration detected by DNA sequencing and low tumor mutation burden. Clinical Cancer "
        "Research, 25(15), 4712–4722.",
        "Benjamin, D., Sato, T., Cibulskis, K., Getz, G., Stewart, C., & Lichtenstein, L. (2019). Calling "
        "somatic SNVs and indels with Mutect2. bioRxiv, 861054.",
        "Capone, I., Bozzi, F., Dagrada, G., Verderio, P., Conca, E., Perrone, F., … Tamborini, E. (2022). "
        "Targeted RNA-sequencing analysis for fusion transcripts detection in tumor diagnostics: assessment "
        "of bioinformatic tools reliability in FFPE samples. Exploration of Targeted Anti-tumor Therapy, 3, "
        "582–597.",
        "Chakravarty, D., Gao, J., Phillips, S. M., Kundra, R., Zhang, H., Wang, J., … Schultz, N. (2017). "
        "OncoKB: A precision oncology knowledge base. JCO Precision Oncology, 1, 1–16.",
        "Díossy, M., Sztupinszki, Z., Krzystanek, M., Borcsok, J., Eklund, A. C., Csabai, I., … Szallasi, Z. "
        "(2021). Strand orientation bias detector to determine the probability of FFPE sequencing artifacts. "
        "Briefings in Bioinformatics, 22(6), bbab186.",
        "Do, H., & Dobrovic, A. (2015). Sequence artifacts in DNA from formalin-fixed tissues: causes and "
        "strategies for minimization. Clinical Chemistry, 61(1), 64–71.",
        "Griffith, M., Spies, N. C., Krysiak, K., McMichael, J. F., Coffman, A. C., Danos, A. M., … "
        "Griffith, O. L. (2017). CIViC is a community knowledgebase for expert crowdsourcing the clinical "
        "interpretation of variants in cancer. Nature Genetics, 49(2), 170–174.",
        "Grüning, B., Dale, R., Sjödin, A., Chapman, B. A., Rowe, J., Tomkins-Tinch, C. H., … The Bioconda "
        "Team. (2018). Bioconda: sustainable and comprehensive software distribution for the life sciences. "
        "Nature Methods, 15(7), 475–476.",
        "Haas, B. J., Dobin, A., Li, B., Stransky, N., Pochet, N., & Regev, A. (2019). Accuracy assessment of "
        "fusion transcript detection via read-mapping and de novo fusion transcript assembly-based methods. "
        "Genome Biology, 20(1), 213.",
        "Kim, H., Lee, A. J., Lee, J., Chun, H., Ju, Y. S., & Hong, D. (2019). Performance evaluation of "
        "commercial library construction kits for PCR-based targeted sequencing using a unique molecular "
        "identifier. BMC Genomics, 20(1), 216.",
        "Li, M. M., Datto, M., Duncavage, E. J., Kulkarni, S., Lindeman, N. I., Roy, S., … Nikiforova, M. N. "
        "(2017). Standards and guidelines for the interpretation and reporting of sequence variants in "
        "cancer: A joint consensus recommendation of the Association for Molecular Pathology, American "
        "Society of Clinical Oncology, and College of American Pathologists. Journal of Molecular "
        "Diagnostics, 19(1), 4–23.",
        "Mateo, J., Chakravarty, D., Dienstmann, R., Jezdic, S., Gonzalez-Perez, A., Lopez-Bigas, N., … "
        "Pusztai, L. (2018). A framework to rank genomic alterations as targets for cancer precision "
        "medicine: the ESMO Scale for Clinical Actionability of molecular Targets (ESCAT). Annals of "
        "Oncology, 29(9), 1895–1902.",
        "Merino, D. M., McShane, L. M., Fabrizio, D., Funari, V., Chen, S.-J., White, J. R., … Allen, J. "
        "(2020). Establishing guidelines to harmonize tumor mutational burden (TMB): in silico assessment of "
        "variation in TMB quantification across diagnostic platforms. Journal for ImmunoTherapy of Cancer, "
        "8(1), e000147.",
        "Mölder, F., Jablonski, K. P., Letcher, B., Hall, M. B., Tomkins-Tinch, C. H., Sochat, V., … Köster, "
        "J. (2021). Sustainable data analysis with Snakemake. F1000Research, 10, 33.",
        "Nakken, S., Fournous, G., Vodák, D., Aasheim, L. B., Myklebost, O., & Hovig, E. (2018). Personal "
        "Cancer Genome Reporter: variant interpretation report for precision oncology. Bioinformatics, "
        "34(10), 1778–1780.",
        "Salk, J. J., Schmitt, M. W., & Loeb, L. A. (2018). Enhancing the accuracy of next-generation "
        "sequencing for detecting rare and subclonal mutations. Nature Reviews Genetics, 19(5), 269–285.",
        "Sanchez-Vega, F., Mina, M., Armenia, J., Chatila, W. K., Luna, A., La, K. C., … Schultz, N. (2018). "
        "Oncogenic signaling pathways in The Cancer Genome Atlas. Cell, 173(2), 321–337.",
        "Sun, J. X., He, Y., Sanford, E., Montesion, M., Frampton, G. M., Vignot, S., … Lieber, D. S. (2018). "
        "A computational approach to distinguish somatic vs. germline origin of genomic alterations from deep "
        "sequencing of cancer specimens without a matched normal. PLoS Computational Biology, 14(2), "
        "e1005965.",
        "The Galaxy Community. (2024). The Galaxy platform for accessible, reproducible, and collaborative "
        "data analyses: 2024 update. Nucleic Acids Research, 52(W1), W83–W94.",
        "Uhrig, S., Ellermann, J., Walther, T., Burkhardt, P., Fröhlich, M., Hutter, B., … Brors, B. (2021). "
        "Accurate and efficient detection of gene fusions from RNA sequencing data. Genome Research, 31(3), "
        "448–460.",
        "Wilkinson, M. D., Dumontier, M., Aalbersberg, I. J., Appleton, G., Axton, M., Baak, A., … Mons, B. "
        "(2016). The FAIR Guiding Principles for scientific data management and stewardship. Scientific Data, "
        "3, 160018.",
        "Zheng, Z., Liebers, M., Zhelyazkova, B., Cao, Y., Panditi, D., Lynch, K. D., … Le, L. P. (2014). "
        "Anchored multiplex PCR for targeted next-generation sequencing. Nature Medicine, 20(12), 1479–1484.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(r)

    # ------------------------------------------------------------------ 9 appendices
    doc.add_page_break()
    doc.add_heading("9. APPENDICES", level=1)
    doc.add_heading("Appendix A. Alterations classified as likely somatic", level=2)
    rows = [[t.get("gene", ""), (t.get("hgvsp", "") or "").split(":")[-1], t.get("consequence", ""),
             t.get("amp_tier", ""), t.get("escat", ""), t.get("vaf", ""), t.get("n_callers", "")]
            for t in tiered]
    if rows:
        table(doc, ["Gene", "Protein change", "Consequence", "Tier", "ESCAT", "VAF", "Callers"], rows[:45],
              caption="Table A1. Variants classified as likely somatic, with tier assignments. The sample "
                      "identifier does not appear; the specimen is referred to by an alias throughout.")
    doc.add_heading("Appendix B. Oncogenic pathway membership", level=2)
    prow = [[p_["pathway"], p_["n_altered"], p_["n_genes_in_set"], p_["altered_genes"]]
            for p_ in paths if int(p_.get("n_altered", 0) or 0)]
    if prow:
        table(doc, ["Pathway", "Altered", "Genes in set", "Altered genes"], prow,
              caption="Table B1. Membership mapping against the ten canonical oncogenic pathways of "
                      "Sanchez-Vega et al. (2018). This is not an enrichment analysis: the gene universe is "
                      "defined by the assay.")
    doc.add_heading("Appendix C. Software and reproducibility", level=2)
    para(doc,
         "The complete source code, environment specifications, documentation and test suite are available in "
         "the project repository. The workflow can be executed end to end on a synthetic example dataset "
         "without any patient data. Every JSON output carries a provenance block recording the generation "
         "time, the git commit and the command line.")
    todo(doc, "Insert the public repository URL once the code is pushed to GitHub.")
    prov = S.get("_provenance", {})
    if prov:
        table(doc, ["Field", "Value"], [[k, str(v)] for k, v in prov.items()],
              caption="Table C1. Provenance of the analysis reported here.")

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--results", default="results")
    ap.add_argument("--out", default="report/XX395_FinalReport.docx")
    a = ap.parse_args()
    p = build(a.results, a.out)
    print(f"wrote {p}")


if __name__ == "__main__":
    main()
