#!/usr/bin/env python3
"""Audit the generated report against the FENS XX395 guidelines, mechanically where possible.

Checks what can be measured from the .docx: fonts, spacing, margins, page numbers, abstract length, section
length against the stated page limits, table captions and whether each table is referred to in the text,
whether every reference is cited, and alphabetical ordering.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt

LIMITS = {  # section number prefix -> page limit from the guidelines
    "2.": 3, "3.1": 1, "3.2": 1, "3.3": 1, "3.4": 3, "4.5": 10, "4.6": 1, "5.": 3, "6.": 1, "7.": 1,
}
LINES_PER_PAGE = 23        # 12 pt Times New Roman, double spaced, 1-inch margins
CHARS_PER_LINE = 95        # 6.5-inch text width


def est_pages(texts):
    lines = sum(max(1, (len(t) // CHARS_PER_LINE) + 1) for t in texts if t.strip())
    return lines / LINES_PER_PAGE


def main(path):
    doc = Document(path)
    ok, bad, warn = [], [], []

    # ---- formatting
    n = doc.styles["Normal"]
    (ok if n.font.name == "Times New Roman" else bad).append(f"body font is {n.font.name} (must be Times New Roman)")
    (ok if n.font.size == Pt(12) else bad).append(f"body size is {n.font.size.pt if n.font.size else '?'} pt (must be 12)")
    (ok if abs(n.paragraph_format.line_spacing - 2.0) < 0.01 else bad).append("body line spacing is double")
    s = doc.sections[0]
    margins = {s.top_margin.inches, s.bottom_margin.inches, s.left_margin.inches, s.right_margin.inches}
    (ok if margins == {1.0} else bad).append(f"margins {sorted(margins)} inch (must all be 1)")
    from lxml import etree
    foot = etree.tostring(s.footer._element).decode()
    (ok if "PAGE" in foot else bad).append("page numbers present in the footer")
    head = "".join(p.text for p in s.header.paragraphs).strip()
    (ok if head else bad).append(f"running head in the page header ({'present' if head else 'MISSING'})")

    # ---- abstract
    paras = doc.paragraphs
    idx = {p.text.strip(): i for i, p in enumerate(paras) if p.style.name.startswith("Heading")}
    def span(start_key, stop_pred):
        i = next((i for i, p in enumerate(paras) if p.text.strip() == start_key), None)
        if i is None:
            return []
        out = []
        for p in paras[i + 1:]:
            if p.style.name.startswith("Heading") and stop_pred(p.text.strip()):
                break
            out.append(p.text)
        return out
    abstract = span("ABSTRACT", lambda t: True)
    words = sum(len(t.split()) for t in abstract if not t.startswith("[TO COMPLETE]"))
    (ok if words <= 250 else bad).append(f"abstract is {words} words (limit 250)")

    # ---- section lengths
    heads = [(i, p.text.strip()) for i, p in enumerate(paras) if p.style.name.startswith("Heading")]
    for pref, limit in LIMITS.items():
        start = next((i for i, t in heads if t.startswith(pref)), None)
        if start is None:
            continue
        nxt = next((i for i, t in heads if i > start and not t.startswith(pref)
                    and len(t.split()[0].rstrip(".")) <= len(pref.rstrip("."))), None)
        body = [p.text for p in paras[start + 1: nxt]]
        pages = est_pages(body)
        msg = f"section {pref} ~{pages:.1f} pages (limit {limit})"
        (ok if pages <= limit else warn).append(
            msg + " -- estimated from character counts; confirm the real pagination in Word")

    # ---- tables: captioned and referred to in the text
    caps = [p.text for p in paras if re.match(r"^Table [A-Z]?\d+\.", p.text.strip())]
    (ok if len(caps) >= len(doc.tables) else bad).append(
        f"{len(caps)} captions for {len(doc.tables)} tables")
    body_text = " ".join(p.text for p in paras if not re.match(r"^Table [A-Z]?\d+\.", p.text.strip()))
    uncited = [c.split(".")[0] for c in caps if c.split(".")[0] not in body_text]
    (ok if not uncited else bad).append(
        f"tables referred to in the text ({'all' if not uncited else 'MISSING: ' + ', '.join(uncited)})")

    # ---- references cited, and alphabetical
    ref_start = next((i for i, t in heads if t.startswith("8.")), None)
    ref_end = next((i for i, t in heads if t.startswith("9.")), len(paras))
    refs = [p.text for p in paras[ref_start + 1: ref_end] if p.text.strip()]
    import unicodedata

    def sortkey(r):
        head = r.split(",")[0].strip() if "," in r.split(".")[0] + "," else r.split(".")[0].strip()
        head = unicodedata.normalize("NFKD", head).encode("ascii", "ignore").decode()
        return head.lower()

    surnames = [sortkey(r) for r in refs]
    (ok if surnames == sorted(surnames) else bad).append("references alphabetised (diacritics ignored)")
    body_all = " ".join(p.text for p in paras[:ref_start])
    import unicodedata as _u
    body_norm = _u.normalize("NFKD", body_all).encode("ascii", "ignore").decode()
    missing = [s_ for s_ in surnames
               if s_.split()[0] not in body_norm.lower() and not s_.startswith("the galaxy")]
    (ok if not missing else bad).append(
        f"every reference cited in the text ({'all' if not missing else 'UNCITED: ' + ', '.join(missing)})")

    # ---- decimals in tables
    over = 0
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for m in re.finditer(r"\d+\.(\d+)", c.text):
                    if len(m.group(1)) > 2:
                        over += 1
    (ok if over == 0 else warn).append(f"{over} table values with more than 2 decimal places")

    todos = sum(1 for p in paras if p.text.startswith("[TO COMPLETE]"))
    warn.append(f"{todos} placeholders still to be filled by the author")

    print(f"AUDIT OF {Path(path).name}\n")
    for label, items in (("PASS", ok), ("FAIL", bad), ("CHECK", warn)):
        for i in items:
            print(f"  [{label:5s}] {i}")
    print(f"\n{len(bad)} failing, {len(warn)} to check, {len(ok)} passing")
    return 1 if bad else 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1] if len(sys.argv) > 1 else "report/BIO395_FinalReport_DRAFT.docx"))
